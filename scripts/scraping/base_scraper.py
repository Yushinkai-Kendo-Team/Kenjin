"""Shared WordPress blog scraping logic for YKC Kenjin.

Provides a SiteConfig dataclass for per-site XPath selectors and filter
rules, plus all the generic scraping/saving/metadata functions that the
site-specific entry-point scripts delegate to.
"""

from __future__ import annotations

import re
import sys
import time
import argparse
from dataclasses import dataclass, field
from pathlib import Path
from urllib.parse import urljoin

# Ensure project src/ is importable
_project_root = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(_project_root / "src"))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

import requests
import yaml
from lxml import html

from kendocenter.config import settings


# ---------------------------------------------------------------------------
# Site configuration
# ---------------------------------------------------------------------------

@dataclass
class SiteConfig:
    """Site-specific selectors and rules for a WordPress blog scraper."""

    site_name: str

    # -- Listing page selectors --
    article_link_xpath: str
    pagination_next_xpath: str

    # -- Article page selectors --
    title_xpath: str
    date_xpath: str

    # -- Optional fields (with defaults) --
    source_folder: str = ""  # Theory subfolder name (e.g., "blogs", "kenshi247")
    category_xpath: str = '//a[@rel="category tag"]'
    content_xpath: str = '//div[contains(@class, "entry-content")]'
    content_fallback_xpath: str = '//div[contains(@class, "post-content")]'
    content_elements_xpath: str = ".//p | .//h2 | .//h3 | .//h4 | .//li"

    # -- Content filtering --
    stop_texts: list[str] = field(default_factory=list)
    skip_prefixes: list[str] = field(default_factory=list)

    # -- Attribution detection (checked against last 5 paragraphs) --
    source_patterns: list[str] = field(default_factory=list)
    translator_patterns: list[str] = field(default_factory=list)
    strip_patterns: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------

def load_urls_yaml(yaml_path: Path) -> dict:
    """Load urls.yaml file."""
    with open(yaml_path, encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def discover_blog_sources(theory_dir: Path, source_filter: str = "") -> list[Path]:
    """Find all subfolders containing urls.yaml (searches recursively)."""
    results = []
    for urls_path in sorted(theory_dir.rglob("urls.yaml")):
        if source_filter:
            rel = urls_path.parent.relative_to(theory_dir).as_posix()
            if rel != source_filter:
                continue
        results.append(urls_path)
    return results


def slugify(text: str) -> str:
    """Convert title to a filename-safe slug."""
    text = text.strip().lower()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text)
    return text[:80].strip("-")


# ---------------------------------------------------------------------------
# Fetching
# ---------------------------------------------------------------------------

def fetch_article_links(
    category_url: str, config: SiteConfig, delay: float = 1.0,
) -> list[dict]:
    """Crawl a WordPress category page and extract all article links.

    Follows pagination (older posts) to get all articles.
    """
    all_links: list[dict] = []
    seen_urls: set[str] = set()
    url: str | None = category_url

    while url:
        print(f"  Fetching: {url}")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 404:
            break
        resp.raise_for_status()
        tree = html.fromstring(resp.text)

        link_els = tree.xpath(config.article_link_xpath)
        for a in link_els:
            href = a.get("href", "")
            title = a.text_content().strip()
            if href and title and href not in seen_urls:
                seen_urls.add(href)
                all_links.append({"url": href, "title": title})

        # Pagination
        older = tree.xpath(config.pagination_next_xpath)
        url = None
        for link in older:
            next_url = link.get("href", "")
            if next_url and next_url not in seen_urls:
                url = next_url
                time.sleep(delay)

    return all_links


def fetch_article_content(url: str, config: SiteConfig) -> dict:
    """Fetch a single blog post and extract its content."""
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    tree = html.fromstring(resp.text)

    # Title
    title_el = tree.xpath(config.title_xpath)
    title = title_el[0].text_content().strip() if title_el else ""

    # Date
    date_el = tree.xpath(config.date_xpath)
    date = ""
    if date_el:
        if isinstance(date_el[0], str):
            date = date_el[0][:10]
        else:
            date = date_el[0].text_content().strip()

    # Categories
    cat_els = tree.xpath(config.category_xpath)
    categories = [c.text_content().strip() for c in cat_els]

    # Main content
    content_el = tree.xpath(config.content_xpath)
    if not content_el:
        content_el = tree.xpath(config.content_fallback_xpath)

    paragraphs: list[str] = []
    if content_el:
        elements = content_el[0].xpath(config.content_elements_xpath)
        for el in elements:
            text = el.text_content().strip()
            if not text:
                continue
            # Stop on sharing/related sections
            if text in config.stop_texts:
                break
            # Skip social/loading text
            if any(text.startswith(p) for p in config.skip_prefixes):
                continue
            paragraphs.append(text)

    # Source / translator attribution
    source_ref = ""
    translator = ""
    for p in paragraphs[-5:]:
        p_lower = p.lower()
        if any(pat in p_lower for pat in config.source_patterns):
            source_ref = p
        if any(pat in p_lower for pat in config.translator_patterns):
            translator = p

    # Remove attribution lines from body
    body_paragraphs = []
    for p in paragraphs:
        p_lower = p.lower()
        if any(pat in p_lower for pat in config.strip_patterns):
            continue
        body_paragraphs.append(p)

    return {
        "url": url,
        "title": title,
        "date": date,
        "categories": categories,
        "paragraphs": body_paragraphs,
        "source_ref": source_ref,
        "translator": translator,
    }


# ---------------------------------------------------------------------------
# Saving
# ---------------------------------------------------------------------------

def save_as_docx(article: dict, output_path: Path) -> None:
    """Save article content as a .docx file."""
    from docx import Document

    doc = Document()
    doc.add_heading(article["title"], level=1)

    if article.get("date"):
        doc.add_paragraph(f"Date: {article['date']}")
    if article.get("source_ref"):
        doc.add_paragraph(f"Source: {article['source_ref']}")
    doc.add_paragraph(f"URL: {article['url']}")
    doc.add_paragraph("")  # Blank separator

    for para in article["paragraphs"]:
        doc.add_paragraph(para)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

def scrape_blog_source(
    urls_yaml_path: Path, config: SiteConfig, dry_run: bool = False,
) -> list[dict]:
    """Scrape all URLs defined in a urls.yaml file."""
    yaml_config = load_urls_yaml(urls_yaml_path)
    source_dir = urls_yaml_path.parent
    scraped: list[dict] = []

    urls_entries = yaml_config.get("urls", [])
    if not urls_entries:
        print(f"  No URLs configured in {urls_yaml_path}")
        return scraped

    delay = yaml_config.get("delay", 1.0)
    default_tags = yaml_config.get("default_tags", [])

    for url_entry in urls_entries:
        url = url_entry if isinstance(url_entry, str) else url_entry.get("url", "")
        url_type = (
            "category" if isinstance(url_entry, str)
            else url_entry.get("type", "category")
        )

        if url_type == "category":
            print(f"\n  Crawling category: {url}")
            links = fetch_article_links(url, config, delay=delay)
            print(f"  Found {len(links)} articles")
        else:
            title = url_entry.get("title", "") if isinstance(url_entry, dict) else ""
            links = [{"url": url, "title": title}]

        for link in links:
            slug = slugify(link["title"]) or slugify(link["url"].split("/")[-2])
            filename = f"{slug}.docx"
            output_path = source_dir / filename

            if output_path.exists():
                print(f"    SKIP (exists): {filename}")
                scraped.append({
                    "filename": filename,
                    "title": link["title"],
                    "url": link["url"],
                    "skipped": True,
                })
                continue

            if dry_run:
                print(f"    WOULD SCRAPE: {link['title'][:60]}")
                print(f"                  -> {filename}")
                continue

            print(f"    Scraping: {link['title'][:60]}...")
            try:
                article = fetch_article_content(link["url"], config)
                if not article["paragraphs"]:
                    print("    WARNING: No content extracted, skipping")
                    continue

                save_as_docx(article, output_path)
                print(f"    Saved: {filename} ({len(article['paragraphs'])} paragraphs)")

                scraped.append({
                    "filename": filename,
                    "title": article["title"],
                    "url": link["url"],
                    "date": article.get("date", ""),
                    "categories": article.get("categories", []),
                    "tags": default_tags + article.get("categories", []),
                    "skipped": False,
                })

                time.sleep(delay)

            except Exception as e:
                print(f"    ERROR: {e}")

    return scraped


def update_metadata_yaml(
    source_dir: Path, scraped: list[dict], config: dict,
) -> None:
    """Update or create metadata.yaml with entries for scraped articles."""
    meta_path = source_dir / "metadata.yaml"

    if meta_path.exists():
        with open(meta_path, encoding="utf-8") as f:
            meta = yaml.safe_load(f) or {}
    else:
        meta = {
            "category": config.get("category", source_dir.name),
            "description": config.get("description", ""),
            "doc_type": config.get("doc_type", "article"),
            "default_language": config.get("default_language", "en"),
        }

    if "files" not in meta:
        meta["files"] = {}

    for article in scraped:
        if article.get("skipped"):
            continue
        filename = article["filename"]
        if filename not in meta["files"]:
            entry: dict = {"title": article.get("title", "")}
            if article.get("url"):
                entry["url"] = article["url"]
            if article.get("date"):
                entry["date"] = article["date"]
            tags = article.get("tags", [])
            if tags:
                entry["tags"] = tags
            meta["files"][filename] = entry

    with open(meta_path, "w", encoding="utf-8") as f:
        yaml.dump(
            meta, f, allow_unicode=True, default_flow_style=False, sort_keys=False,
        )

    print(f"\n  Updated: {meta_path}")


def run_scraper(
    config: SiteConfig,
    source_filter: str = "",
    dry_run: bool = False,
) -> None:
    """Main entry point: discover sources, scrape, update metadata."""
    theory_dir = settings.theory_path
    print(f"=== YKC Kenjin Blog Scraper ({config.site_name}) ===")
    print(f"Source directory: {theory_dir}")

    # Use config.source_folder as default filter so each scraper
    # only processes its own Theory/ subfolder.
    effective_filter = source_filter or config.source_folder
    urls_files = discover_blog_sources(theory_dir, effective_filter)
    if not urls_files:
        print("No urls.yaml files found in source directories.")
        print("Create a urls.yaml in a KENDO_THEORY_DIR subfolder. See docs/adding-content.md")
        sys.exit(1)

    for urls_path in urls_files:
        source_dir = urls_path.parent
        print(f"\n--- Processing: {source_dir.name}/ ---")

        yaml_config = load_urls_yaml(urls_path)
        scraped = scrape_blog_source(urls_path, config, dry_run=dry_run)

        if not dry_run and scraped:
            new_articles = [a for a in scraped if not a.get("skipped")]
            if new_articles:
                update_metadata_yaml(source_dir, scraped, yaml_config)
                print(f"  New articles: {len(new_articles)}")
            else:
                print("  No new articles to add.")

    if not dry_run:
        print("\n=== Scraping Complete ===")
        print("Next steps:")
        print("  1. Review the scraped .docx files")
        print("  2. Run ingestion: .venv/Scripts/python.exe scripts/ingest_all.py --reset")
        print("  3. Verify: .venv/Scripts/python.exe scripts/verify_pipeline.py")


def run_scraper_all(source_filter: str = "", dry_run: bool = False) -> None:
    """Run all configured blog scrapers by discovering urls.yaml files.

    Each urls.yaml is matched to a SiteConfig based on the folder or
    config content.  Falls back to a generic WordPress config.
    """
    from scrape_kendo3ka import KENDO3KA_CONFIG
    from scrape_kenshi247 import KENSHI247_CONFIG
    from scrape_kendoinfo import KENDOINFO_CONFIG
    from scrape_nanseikan import NANSEIKAN_CONFIG
    from scrape_kendophilosophy import KENDOPHILOSOPHY_CONFIG

    SITE_CONFIGS = {
        "blogs/kendo3ka": KENDO3KA_CONFIG,
        "blogs/kenshi247": KENSHI247_CONFIG,
        "blogs/kendoinfo": KENDOINFO_CONFIG,
        "blogs/nanseikan": NANSEIKAN_CONFIG,
        "blogs/kendophilosophy": KENDOPHILOSOPHY_CONFIG,
    }

    theory_dir = settings.theory_path
    urls_files = discover_blog_sources(theory_dir, source_filter)

    for urls_path in urls_files:
        rel = urls_path.parent.relative_to(theory_dir).as_posix()
        config = SITE_CONFIGS.get(rel)
        if config is None:
            print(f"WARNING: No scraper config for '{rel}/', skipping")
            continue
        run_scraper(config, source_filter=rel, dry_run=dry_run)


def make_arg_parser(description: str) -> argparse.ArgumentParser:
    """Create a standard CLI argument parser for scraper entry points."""
    parser = argparse.ArgumentParser(description=description)
    parser.add_argument("--dry-run", action="store_true", help="Show what would be scraped")
    parser.add_argument("--source", default="", help="Scrape specific subfolder only")
    return parser
