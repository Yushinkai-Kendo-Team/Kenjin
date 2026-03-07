"""Parse .docx articles about kendo into structured documents.

The articles may contain:
- English text (or English translation)
- Vietnamese translation in the second half
- Rich kendo terminology used in context
- Biographical information about sensei

The parser extracts metadata (title, date, source, translator) and
splits the content into English and Vietnamese sections.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


@dataclass
class ArticleDocument:
    """A parsed article with metadata and content."""

    filename: str
    title: str = ""
    date: str = ""
    source_publication: str = ""
    translator: str = ""
    subject_name: str = ""
    subject_rank: str = ""

    # Content split by language
    english_paragraphs: list[str] = field(default_factory=list)
    vietnamese_paragraphs: list[str] = field(default_factory=list)

    @property
    def english_text(self) -> str:
        return "\n\n".join(self.english_paragraphs)

    @property
    def vietnamese_text(self) -> str:
        return "\n\n".join(self.vietnamese_paragraphs)


# Vietnamese-specific characters for language detection
VIETNAMESE_CHARS = set("ệờạướồểữịựợặắảẩẫấầẻẽỏốổỗộớợừửữỳỵỷỹđĐ")


def _is_vietnamese(text: str) -> bool:
    """Check if text contains Vietnamese-specific characters."""
    return any(c in VIETNAMESE_CHARS for c in text) and len(text) > 20


def _extract_metadata(paragraphs: list[str], filename: str) -> dict:
    """Extract article metadata from the first few paragraphs."""
    metadata = {
        "title": "",
        "date": "",
        "source_publication": "",
        "translator": "",
        "subject_name": "",
        "subject_rank": "",
    }

    for i, text in enumerate(paragraphs[:10]):
        text = text.strip()
        if not text:
            continue

        # Title is usually the first non-empty paragraph
        if not metadata["title"] and i < 3:
            metadata["title"] = text

        # Date patterns
        date_match = re.match(r"^(\d{2}/\d{2}/\d{4})", text)
        if date_match:
            metadata["date"] = date_match.group(1)

        # Source publication detection (matches lines with a year.issue pattern
        # like "Magazine Name 2025.1" — common in kendo magazine articles)
        issue_match = re.search(r"(\d{4}\.\d+)", text)
        if issue_match:
            metadata["source_publication"] = text.strip()
            metadata["date"] = issue_match.group(1)

        # Translator
        if "translation" in text.lower() or "dịch:" in text.lower():
            metadata["translator"] = text

        # Try to extract subject name from filename (e.g., "Sensei-Name.docx")
        # Matches capitalized two-word names separated by hyphen or space
        name_match = re.search(r"([A-Z][a-z]+[-. ][A-Z][a-z]+)", filename)
        if name_match:
            metadata["subject_name"] = name_match.group(1).replace("-", " ")

    # Extract subject name from title if not from filename
    if not metadata["subject_name"] and metadata["title"]:
        # Common patterns: "Title (Name)" or "Title – Name"
        name_in_title = re.search(
            r"[–\-]\s*([A-Z][a-z]+ [A-Z][a-z]+)\s*$", metadata["title"]
        )
        if name_in_title:
            metadata["subject_name"] = name_in_title.group(1)
        else:
            paren_name = re.search(r"\(([A-Z][a-z]+ [A-Z][a-z]+)\)", metadata["title"])
            if paren_name:
                metadata["subject_name"] = paren_name.group(1)

    return metadata


def parse_docx(docx_path: str | Path) -> ArticleDocument:
    """Parse a .docx article into structured content.

    Args:
        docx_path: Path to the .docx file

    Returns:
        ArticleDocument with metadata and split English/Vietnamese content.
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))

    # Collect all paragraph texts
    all_paragraphs = [p.text.strip() for p in doc.paragraphs]

    # Extract metadata from early paragraphs
    metadata = _extract_metadata(all_paragraphs, docx_path.name)

    # Find the boundary between English and Vietnamese
    vietnamese_start = None
    for i, text in enumerate(all_paragraphs):
        if _is_vietnamese(text):
            # Look back a few lines for the Vietnamese section header
            vietnamese_start = i
            # Check if previous lines are headers/markers for the Vietnamese section
            for j in range(max(0, i - 3), i):
                prev = all_paragraphs[j].strip()
                if prev and (
                    "dịch" in prev.lower()
                    or "việt" in prev.lower()
                    or "bản dịch" in prev.lower()
                    or _is_vietnamese(prev)
                ):
                    vietnamese_start = j
                    break
            break

    # Split into English and Vietnamese sections
    if vietnamese_start is not None:
        english_raw = all_paragraphs[:vietnamese_start]
        vietnamese_raw = all_paragraphs[vietnamese_start:]
    else:
        english_raw = all_paragraphs
        vietnamese_raw = []

    # Filter out empty paragraphs and metadata lines for content
    english_paragraphs = [
        p for p in english_raw if p.strip() and len(p.strip()) > 5
    ]
    vietnamese_paragraphs = [
        p for p in vietnamese_raw if p.strip() and len(p.strip()) > 5
    ]

    # Handle Vietnamese-only documents: if the "English" section is just a
    # title or metadata (very little real content), treat the whole document
    # as Vietnamese-only.  The title/metadata lines get prepended to the
    # Vietnamese section so nothing is lost.
    en_total_chars = sum(len(p) for p in english_paragraphs)
    if vietnamese_paragraphs and en_total_chars < 300:
        # Move any short "English" lines into the Vietnamese section
        vietnamese_paragraphs = english_paragraphs + vietnamese_paragraphs
        english_paragraphs = []

    return ArticleDocument(
        filename=docx_path.name,
        title=metadata["title"],
        date=metadata["date"],
        source_publication=metadata["source_publication"],
        translator=metadata["translator"],
        subject_name=metadata["subject_name"],
        english_paragraphs=english_paragraphs,
        vietnamese_paragraphs=vietnamese_paragraphs,
    )


def parse_all_docx(directory: str | Path) -> list[ArticleDocument]:
    """Parse all .docx files in a directory.

    Args:
        directory: Path to directory containing .docx files

    Returns:
        List of ArticleDocument objects.
    """
    directory = Path(directory)
    articles = []
    for docx_file in sorted(directory.glob("*.docx")):
        article = parse_docx(docx_file)
        articles.append(article)
    return articles


if __name__ == "__main__":
    import sys

    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    from kendocenter.config import settings

    theory_dir = settings.theory_path
    articles = parse_all_docx(theory_dir)

    for article in articles:
        print(f"=== {article.filename} ===")
        print(f"  Title: {article.title}")
        print(f"  Date: {article.date}")
        print(f"  Source: {article.source_publication}")
        print(f"  Translator: {article.translator}")
        print(f"  Subject: {article.subject_name}")
        print(f"  English paragraphs: {len(article.english_paragraphs)}")
        print(f"  Vietnamese paragraphs: {len(article.vietnamese_paragraphs)}")
        if article.english_paragraphs:
            print(f"  First EN para: {article.english_paragraphs[0][:100]}...")
        if article.vietnamese_paragraphs:
            print(f"  First VN para: {article.vietnamese_paragraphs[0][:100]}...")
        print()
